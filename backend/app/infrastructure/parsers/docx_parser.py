"""
Python-docx parser for Microsoft Word (.docx) documents.
Extracts headings, paragraphs, and tables.
"""
from typing import Optional, Dict, Any
import io
import docx
from app.domain.models import Document, DocumentMetadata
from app.core.logging import logger


class DOCXParser:
    """Extracts paragraphs and tables from DOCX files."""

    def parse(self, file_content: bytes, filename: str, metadata: Optional[Dict[str, Any]] = None) -> Document:
        custom_meta = metadata or {}
        doc_stream = io.BytesIO(file_content)
        doc = docx.Document(doc_stream)

        text_blocks = []
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                if paragraph.style.name.startswith("Heading"):
                    text_blocks.append(f"## {text}")
                else:
                    text_blocks.append(text)

        # Process tables
        for table in doc.tables:
            table_rows = []
            for row in table.rows:
                row_cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                if any(row_cells):
                    table_rows.append(" | ".join(row_cells))
            if table_rows:
                text_blocks.append("\n".join(table_rows))

        full_content = "\n\n".join(text_blocks)
        core_props = doc.core_properties
        title = core_props.title or filename
        author = core_props.author

        logger.info(f"Parsed DOCX '{filename}': {len(text_blocks)} blocks, {len(full_content)} chars.")

        return Document(
            filename=filename,
            content=full_content,
            metadata=DocumentMetadata(
                title=title,
                author=author,
                file_type="docx",
                file_size=len(file_content),
                custom_metadata=custom_meta
            )
        )
